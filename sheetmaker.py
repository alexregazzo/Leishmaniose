import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import registro
import typing
import utils
import settings


def _make_sheet(*, filename, regs: typing.List[registro.Registro]) -> typing.List[str]:
    if len(regs) > 18:
        results = []
        for i in range(0, len(regs), 18):
            fn, fe = os.path.splitext(filename)
            results.extend(_make_sheet(filename=(fn + f"_F{int(i // 18)}" + fe), regs=regs[i:i + 18]))
        return results
    doc = docx.Document(os.path.join(settings.ROOT_DIRPATH, "planilha_relatorio.docx"))
    dados = utils.get_data()
    table = doc.tables[0]
    estatistica = {"teste": {}}
    for i, reg in enumerate(regs):
        estatistica["teste"]["coletado"] = estatistica["teste"].get("coletado", 0) + (1 if reg["reg_teste_resultado"] in ["1", "2"] else 0)
        estatistica["teste"]["positivo"] = estatistica["teste"].get("positivo", 0) + (1 if reg["reg_teste_resultado"] == "2" else 0)
        estatistica["teste"]["negativo"] = estatistica["teste"].get("negativo", 0) + (1 if reg["reg_teste_resultado"] == "1" else 0)
        for dado in dados:
            if dado["planilha"] is None:
                continue
            texto = reg[f"reg_{dado['nome']}"]
            if dado["tipo"] == "date":
                texto = utils.strstd2strusr(texto)
            if dado["valores"] is not None:
                if texto == "0":
                    texto = ""
            table.rows[3 + i].cells[dado['planilha']].text = texto
            table.rows[3 + i].cells[dado['planilha']].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[21].cells[6].text = str(estatistica["teste"].get("coletado", 0))
    table.rows[22].cells[6].text = str(estatistica["teste"].get("positivo", 0))
    table.rows[23].cells[6].text = str(estatistica["teste"].get("negativo", 0))
    table.rows[21].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[22].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[23].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if os.path.splitext(filename)[1] != ".docx":
        filename += ".docx"
    filepath = utils.make_unique_filepath(os.path.join(settings.REPORT_DIRPATH, filename))
    logger.debug(f"Report file save as: {filepath}")
    doc.save(filepath)
    return [filepath]


def make_filename(*, quadra: str) -> str:
    return utils.get_datetime_for_file() + "_Q" + quadra.zfill(4)


def make_sheet(*, quadra: typing.Union[int, str] = None, regs: typing.List[registro.Registro] = None) -> typing.List[str]:
    logger.debug(str(quadra) + str(regs))
    if quadra is not None:
        quadra = str(quadra)
    if regs is None:
        regs = registro.Registro.get_all()
    results = []
    if quadra is None:
        logger.debug("No block")
        split_regs = {}
        for reg in regs:
            split_regs[reg["reg_quadra"]] = split_regs.get(reg["reg_quadra"], []) + [reg]
        for quadra in split_regs.keys():
            fn = make_filename(quadra=quadra)
            logger.debug(f"Filename: {fn}")
            results.extend(_make_sheet(filename=fn, regs=split_regs[quadra]))
    else:
        # quadra is not none
        logger.debug("Block defined")
        regs = [reg for reg in regs if reg["reg_quadra"] == quadra]
        fn = make_filename(quadra=quadra)
        logger.debug(f"Filename: {fn}")
        results.extend(_make_sheet(filename=fn, regs=regs))
    return results


logger = utils.get_logger(__file__)
if __name__ == "__main__":
    pass
