import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import registro
from datetime import datetime
import typing
import utils
from zipfile import ZipFile


def make_unique_filepath(fpath: str) -> str:
    if os.path.exists(fpath):
        fname, fext = os.path.splitext(fpath)
        fpath = fname + " ({})" + fext
        k = 0
        while os.path.exists(fpath.format(k)):
            k += 1
        fpath = fpath.format(k)
    return fpath


def parse_date(datestr: str) -> str:
    try:
        d = datetime.strptime(datestr, "%Y-%m-%d")
    except ValueError:
        return datestr
    else:
        return d.strftime("%d/%m/%Y")


def _make_sheet(*, filename, regs: typing.List[registro.Registro]) -> typing.List[str]:
    if len(regs) > 18:
        results = []
        for i in range(0, len(regs), 18):
            fn, fe = os.path.splitext(filename)
            results.extend(_make_sheet(filename=(fn + f"_F{int(i // 18)}" + fe), regs=regs[i:i + 18]))
        return results
    doc = docx.Document("base.docx")
    dados = utils.get_data()
    table = doc.tables[0]
    estatistica = {"teste": {}}
    for i, reg in enumerate(regs):
        estatistica["teste"]["coletado"] = estatistica["teste"].get("coletado", 0) + (1 if reg["reg_teste_resultado"] in ["1", "2"] else 0)
        estatistica["teste"]["positivo"] = estatistica["teste"].get("positivo", 0) + (1 if reg["reg_teste_resultado"] == "2" else 0)
        estatistica["teste"]["negativo"] = estatistica["teste"].get("negativo", 0) + (1 if reg["reg_teste_resultado"] == "1" else 0)
        for dado in dados:
            if dado["planilha"] is None:
                continue
            texto = reg[f"reg_{dado['nome']}"]
            if dado["tipo"] == "date":
                texto = parse_date(texto)
            if dado["valores"] is not None:
                if texto == "0":
                    texto = ""
            table.rows[3 + i].cells[dado['planilha']].text = texto
            table.rows[3 + i].cells[dado['planilha']].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[21].cells[6].text = str(estatistica["teste"].get("coletado", 0))
    table.rows[22].cells[6].text = str(estatistica["teste"].get("positivo", 0))
    table.rows[23].cells[6].text = str(estatistica["teste"].get("negativo", 0))
    table.rows[21].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[22].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[23].cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    root_dir = "relatorios"
    os.makedirs(root_dir, exist_ok=True)
    if os.path.splitext(filename)[1] != ".docx":
        filename += ".docx"
    filepath = make_unique_filepath(os.path.join(root_dir, filename))
    doc.save(filepath)
    return [filepath]


def get_datetime_for_file() -> str:
    return datetime.now().strftime("%Y-%m-%d_%H-%M-%S")


def make_filename(*, quadra: str) -> str:
    return get_datetime_for_file() + "_Q" + quadra.zfill(4)


def make_sheet(*, quadra: typing.Union[int, str] = None, regs: typing.List[registro.Registro] = None) -> typing.List[str]:
    if quadra is not None:
        quadra = str(quadra)
    if regs is None:
        regs = registro.Registro.get_all()
    results = []
    if quadra is None:
        split_regs = {}
        for reg in regs:
            split_regs[reg["reg_quadra"]] = split_regs.get(reg["reg_quadra"], []) + [reg]
        for quadra in split_regs.keys():
            fn = make_filename(quadra=quadra)
            results.extend(_make_sheet(filename=fn, regs=split_regs[quadra]))
    else:
        # quadra is not none
        regs = [reg for reg in regs if reg["reg_quadra"] == quadra]
        fn = make_filename(quadra=quadra)
        results.extend(_make_sheet(filename=fn, regs=regs))
    return results


def zip_files(files: typing.List[str], *, remove_dirs: bool = False) -> str:
    root_dir = "zips"
    os.makedirs(root_dir, exist_ok=True)
    zippath = make_unique_filepath(os.path.join(root_dir, get_datetime_for_file() + '.zip'))
    with ZipFile(zippath, 'w') as myzip:
        for file in files:
            myzip.write(file, arcname=(os.path.split(file)[1] if remove_dirs else None))
    return zippath


if __name__ == "__main__":
    pass
    # regs = registro.Registro.get_all() * 3
    # print(_make_sheet(filename=make_filename(quadra="10"), regs=regs))
